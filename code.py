# math_question_generator.py
import random
from docx import Document
from docx.shared import Pt

def generate_uniform_question():
    """Generates uniform combination question"""
    shirt_colors = ['Blue', 'Red', 'White', 'Green', 'Black', 'Yellow']
    pants_colors = ['Black', 'Khaki', 'Navy', 'Gray', 'Blue']
    
    # Randomly select 3-4 shirts and 2-3 pants
    selected_shirts = random.sample(shirt_colors, k=random.randint(3,4))
    selected_pants = random.sample(pants_colors, k=random.randint(2,3))
    
    # Create question text
    question = "Each student at Lincoln High School wears a uniform consisting of 1 shirt and 1 pair of pants. "
    question += "The table shows the colors available for each item of clothing. How many different uniforms are possible?\n\n"
    
    # Add table
    question += "| Shirt Color | Pants Color |\n"
    question += "|-------------|-------------|\n"
    for shirt in selected_shirts:
        if selected_pants:  # Ensure we don't pop from empty list
            pants = selected_pants.pop(0) if random.random() > 0.2 else ""
            question += f"| {shirt} | {pants} |\n"
        else:
            question += f"| {shirt} | |\n"
    
    # Calculate correct answer
    num_shirts = len(selected_shirts)
    num_pants = len([p for p in selected_pants if p]) + (len(selected_shirts) - len(selected_pants))
    correct_answer = num_shirts * num_pants
    
    # Generate options
    options = [
        correct_answer - 2,
        correct_answer - 1,
        correct_answer,
        correct_answer + 1,
        correct_answer + 2
    ]
    options = list(set(options))  # Remove duplicates
    random.shuffle(options)
    
    return {
        "question": question,
        "options": options,
        "correct_answer": correct_answer,
        "explanation": f"There are {num_shirts} shirt options and {num_pants} pants options, so there are {num_shirts} × {num_pants} = {correct_answer} possible uniform combinations."
    }

def save_to_word(questions):
    """Saves questions to Word document"""
    doc = Document()
    
    # Add title
    doc.add_heading('Generated Math Questions', level=1)
    
    for i, q in enumerate(questions, 1):
        # Add question
        doc.add_paragraph(f"Question {i}", style='Heading 2')
        doc.add_paragraph(q['question'])
        
        # Add options
        for j, option in enumerate(q['options'], 1):
            doc.add_paragraph(f"{chr(64+j)}) {option}")
        
        # Add explanation
        doc.add_paragraph("Explanation:", style='Heading 3')
        doc.add_paragraph(q['explanation'])
    
    doc.save('math_questions.docx')

if __name__ == "__main__":
    # Generate 2 questions
    questions = [generate_uniform_question() for _ in range(2)]
    save_to_word(questions)
    print("Questions generated and saved to math_questions.docx")
