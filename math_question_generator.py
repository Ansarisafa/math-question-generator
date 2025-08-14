import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
import os

def generate_uniform_question():
    """Generates uniform combination question"""
    return {
        "question": (
            "Each student at Central Middle School wears a uniform consisting of "
            "1 shirt and 1 pair of pants. The table shows the colors available for "
            "each item of clothing. How many different uniforms are possible?\n\n"
            "| Shirt Color | Pants Color |\n"
            "| :---: | :---: |\n"
            "| Tan | Black |\n"
            "| Red | Khaki |\n"
            "| White | Navy |\n"
            "| Yellow | |\n"
        ),
        "options": {
            "A": "Three",
            "B": "Four",
            "C": "Seven",
            "D": "Ten",
            "E": "Twelve"
        },
        "correct_answer": "E",
        "explanation": (
            r"Multiply the number of shirt options (4) by pants options (3). "
            r"The empty cell implies all shirts can pair with all pants: \(4 \times 3 = 12\) combinations."
        ),
        "subject": "Quantitative Math",
        "unit": "Problem Solving",
        "topic": "Data Analysis",
        "image": None
    }

def generate_ball_packing():
    """Generates ball packing question with diagram"""
    # Create diagram
    fig, ax = plt.subplots(figsize=(4, 3))
    for x in [0, 1, 2]:  # 3 columns
        for y in [0, 1]:  # 2 rows
            ax.add_patch(plt.Circle((x*4+2, y*4+2), 2, fill=False, linewidth=2))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_aspect('equal')
    plt.savefig('ball_packing.png', dpi=120, bbox_inches='tight')
    plt.close()
    
    return {
        "question": "The top view of a rectangular package of 6 tightly packed balls is shown below.",
        "options": {
            "A": r"\(2 \times 3 \times 6\)",
            "B": r"\(4 \times 6 \times 6\)",
            "C": r"\(2 \times 4 \times 6\)",
            "D": r"\(4 \times 8 \times 12\)",
            "E": r"\(6 \times 8 \times 12\)"
        },
        "correct_answer": "B",
        "explanation": (
            r"Each ball has diameter \(4\) cm. For 3×2 arrangement:"
            "\n- Width: \(3 \times 4 = 12\) cm"
            "\n- Height: \(2 \times 4 = 8\) cm"
            "\n- Depth: \(4\) cm"
        ),
        "subject": "Quantitative Math",
        "unit": "Problem Solving",
        "topic": "Geometry",
        "image": "ball_packing.png"
    }

def create_document():
    """Creates properly formatted Word document"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    questions = [generate_uniform_question(), generate_ball_packing()]
    
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'Question {i}', level=1)
        doc.add_paragraph(q['question'])
        
        # Insert image if exists
        if q['image'] and os.path.exists(q['image']):
            doc.add_picture(q['image'], width=Inches(3.0))
        
        # Add options
        for opt, text in q['options'].items():
            doc.add_paragraph(f"{opt}) {text}", style='List Bullet')
        
        # Add metadata
        meta = doc.add_paragraph()
        meta.add_run("Curriculum Tags:\n").bold = True
        meta.add_run(f"{q['subject']} > {q['unit']} > {q['topic']}")
        
        # Explanation
        doc.add_heading('Explanation', level=2)
        doc.add_paragraph(q['explanation'])
        
        if i < len(questions):
            doc.add_page_break()
    
    doc.save('math_assessment.docx')
    if os.path.exists('ball_packing.png'):
        os.remove('ball_packing.png')

if __name__ == "__main__":
    create_document()
    print("Assessment generated: math_assessment.docx")
